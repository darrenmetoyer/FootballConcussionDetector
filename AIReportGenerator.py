import json
import os
import uuid
from docx import Document
from google import genai
from dotenv import load_dotenv

load_dotenv("APIKeys.env")

generatedReportsFolder = "GeneratedReports"
os.makedirs(generatedReportsFolder, exist_ok=True)

geminiClient = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))


def buildReportParagraph(detectorResult):
    detectorJson = json.dumps(detectorResult, indent=2)

    prompt = f"""
You are writing a professional concussion analysis summary.

Tell the user at what frame the impact occurred and the second of the impact (30 frames per second).

Also, give them the 4 coordinates of the bounding box around the head.

Use only the facts in the JSON below.
Do not invent missing details.
Write exactly one concise paragraph suitable for a report.

Detector JSON:
{detectorJson}
"""

    response = geminiClient.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
    )

    return response.text.strip()


def createReportDoc(videoName, reportParagraph):
    docId = str(uuid.uuid4())
    filePath = os.path.join(generatedReportsFolder, f"{docId}.docx")

    document = Document()
    document.add_heading("Concussion Analysis Report", level=1)
    document.add_paragraph(f"Video: {videoName}")
    document.add_paragraph(f"Generated: {__import__('datetime').datetime.now().strftime('%Y-%m-%d')}")
    document.add_paragraph("")
    document.add_paragraph(reportParagraph)
    document.save(filePath)

    return docId, filePath